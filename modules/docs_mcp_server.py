#!/usr/bin/env python3
"""
modules/docs_mcp_server.py — MCP server exposing docx/xlsx create + read as
direct tool calls, wrapping python-docx / openpyxl (already in this repo's
venv) instead of running ad-hoc scripts through Bash.

Built in-house instead of adopting a third-party MCP server after auditing
the two most relevant options (GongRzhe/Office-Word-MCP-Server,
haris-musa/excel-mcp-server): the Word one is archived with an open,
unpatched "no path restriction on file operations" issue; the Excel one's
own path sandbox explicitly does not apply to stdio transport (the mode
used here), per its open issue #131, and its write_data_to_excel path has
no formula-injection guard. This server closes both gaps from the start:

  1. Writes are confined to an allowlist of roots (ALLOWED_WRITE_ROOTS) —
     WORKSPACE_DIR (the on-demand bot's scratch dir) and this repo's
     data/documents/ — checked against the *resolved* real path, plus a
     required file extension per tool, so nothing here can create or
     overwrite an arbitrary file elsewhere on disk (e.g. .env, .py).
  2. write_xlsx sanitizes any plain data cell whose string value starts
     with =, +, -, or @ (prefixing a literal apostrophe) so it can never
     be auto-promoted to a live formula by openpyxl — the exact class of
     bug the audit found unmitigated in write_data_to_excel upstream.
     Real formulas are only ever written through the explicit `formulas`
     parameter, never through plain data values.

Reads are not path-restricted: python-docx/openpyxl require the target to
parse as a valid OOXML zip, which rules out the arbitrary-local-file-read
class of bug the Word server's add_picture tool had (any bytes handed to
Document()/load_workbook() that aren't real docx/xlsx just fail to parse).

Formatting conventions (professional fonts, financial-model color coding,
number formats) are ported from Anthropic's official docx/xlsx skills
(~/.claude/plugins/marketplaces/anthropic-agent-skills/skills/{docx,xlsx}/
SKILL.md) so output quality tracks what Claude.ai itself produces, without
adopting that skill's actual mechanism (docx-js/pandoc/LibreOffice via
Bash) — everything here stays a single structured tool call. The one thing
NOT replicated: LibreOffice-verified formula recalculation. openpyxl writes
formula strings with no cached value, so a formula cell written here reads
back as None via read_xlsx (or any other data_only reader) until the file
is opened once in Excel/LibreOffice — there's no way around that without
LibreOffice actually installed.
"""
import os
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from dotenv import load_dotenv
load_dotenv(ROOT / ".env")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import openpyxl
from openpyxl.styles import Font, PatternFill

from mcp.server import MCPServer

WORKSPACE_DIR = Path("/home/shanvi/claude_agent")
REPO_DOCS_DIR = ROOT / "data" / "documents"
ALLOWED_WRITE_ROOTS = [WORKSPACE_DIR, REPO_DOCS_DIR]

_FORMULA_TRIGGER_CHARS = ("=", "+", "-", "@")

# Financial-model color conventions from the official xlsx skill: blue for
# hardcoded inputs, black for formulas, green for same-workbook cross-sheet
# links, red for links to another file, yellow fill for key assumptions.
_FINANCIAL_CELL_STYLES = {
    "input":          {"color": "0000FF"},
    "formula":        {"color": "000000"},
    "link":           {"color": "008000"},
    "external_link":  {"color": "FF0000"},
    "assumption":     {"fill": "FFFF00"},
}

_HEADER_FILL = "D9D9D9"  # light gray, used for both docx table headers and xlsx header rows

mcp = MCPServer("docs")


def _validate_write_path(path: str, required_ext: str) -> Path:
    """Resolve `path`, require it end in `required_ext`, and require it fall
    under one of ALLOWED_WRITE_ROOTS. Raises ValueError otherwise. Creates
    the parent directory (still inside the allowed root) if needed."""
    if not path.lower().endswith(required_ext):
        raise ValueError(f"Path must end with {required_ext}: {path}")

    resolved = Path(os.path.realpath(path))
    for root in ALLOWED_WRITE_ROOTS:
        root_resolved = Path(os.path.realpath(root))
        if resolved == root_resolved or root_resolved in resolved.parents:
            resolved.parent.mkdir(parents=True, exist_ok=True)
            return resolved

    allowed = ", ".join(str(r) for r in ALLOWED_WRITE_ROOTS)
    raise ValueError(f"Refusing to write outside allowed directories ({allowed}): {path}")


def _sanitize_cell(value):
    """Defuse Excel formula-injection: a plain data string starting with
    =, +, -, or @ gets a literal leading apostrophe so openpyxl (and Excel)
    treat it as text, never as a live formula. Non-strings pass through."""
    if isinstance(value, str) and value.startswith(_FORMULA_TRIGGER_CHARS):
        return "'" + value
    return value


# ─── docx helpers ──────────────────────────────────────────────────────────

def _shade_table_cell(cell, hex_color: str) -> None:
    """Set a table cell's background fill via raw OOXML (python-docx has no
    high-level API for this)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_paragraph(doc, para_spec, font: str) -> None:
    """`para_spec` is either a plain string, or a dict {"runs": [{"text":,
    "bold":, "italic":, "color": "RRGGBB", "size": pt_int}, ...]} for
    mixed inline formatting within one paragraph."""
    p = doc.add_paragraph()
    if isinstance(para_spec, str):
        run = p.add_run(para_spec)
        run.font.name = font
    else:
        for r in para_spec.get("runs", []):
            run = p.add_run(r.get("text", ""))
            run.font.name = font
            if r.get("bold"):
                run.bold = True
            if r.get("italic"):
                run.italic = True
            if r.get("size"):
                run.font.size = Pt(int(r["size"]))
            if r.get("color"):
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor.from_string(r["color"])


@mcp.tool()
def write_docx(
    path: str,
    title: str | None = None,
    sections: list[dict] | None = None,
    tables: list[dict] | None = None,
    font: str = "Calibri",
) -> dict:
    """Create a formatted .docx file. `path` must end in .docx and resolve
    under /home/shanvi/claude_agent or data/documents/ in this repo.

    `title` becomes the document's Title-styled heading (optional).
    `font` sets the document-wide default font (Normal style) — use a
    professional font (Calibri, Arial, Times New Roman) unless told otherwise.

    `sections`: list of {"heading": str|None, "level": int (1-4, default 1),
    "paragraphs": list}. Each paragraph is either a plain string, or a dict
    {"runs": [{"text": str, "bold": bool, "italic": bool, "color": "RRGGBB",
    "size": int}, ...]} for mixed inline formatting (e.g. a bolded lead-in
    followed by plain text in the same paragraph). Omit `sections` (or
    leave "heading" None) for plain prose with no headings — e.g. a story.

    `tables`: list of {"heading": str|None, "header_row": list[str]|None,
    "rows": list[list[str]]} — each becomes an optional heading, an
    optional bold+shaded header row, then the data rows.
    """
    try:
        dest = _validate_write_path(path, ".docx")
    except ValueError as exc:
        return {"written": False, "error": str(exc)}

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.name = font

    for section in sections or []:
        heading = section.get("heading")
        if heading:
            h = doc.add_heading(heading, level=int(section.get("level", 1)))
            for run in h.runs:
                run.font.name = font
        for para in section.get("paragraphs", []):
            _add_paragraph(doc, para, font)

    for table in tables or []:
        heading = table.get("heading")
        if heading:
            h = doc.add_heading(heading, level=int(table.get("level", 2)))
            for run in h.runs:
                run.font.name = font
        header_row = table.get("header_row")
        rows = table.get("rows", [])
        ncols = len(header_row) if header_row else (len(rows[0]) if rows else 0)
        if ncols == 0:
            continue
        t = doc.add_table(rows=0, cols=ncols)
        t.style = "Table Grid"
        if header_row:
            cells = t.add_row().cells
            for i, val in enumerate(header_row):
                cells[i].text = str(val)
                _shade_table_cell(cells[i], _HEADER_FILL)
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.name = font
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.font.name = font

    doc.save(str(dest))
    return {"written": True, "path": str(dest)}


@mcp.tool()
def read_docx(path: str) -> dict:
    """Read a .docx file's content back out: an ordered list of blocks
    (headings and paragraphs, in document order) plus a separate list of
    tables (each a list of row lists). No path restriction on reads — an
    invalid/non-docx file simply fails to parse rather than leaking bytes.
    """
    try:
        doc = Document(path)
    except Exception as exc:
        return {"error": str(exc)}

    blocks = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading") or style == "Title":
            level = 0 if style == "Title" else int(style.replace("Heading ", "") or 1)
            blocks.append({"type": "heading", "level": level, "text": p.text})
        else:
            blocks.append({"type": "paragraph", "text": p.text})

    tables = [
        [[cell.text for cell in row.cells] for row in t.rows]
        for t in doc.tables
    ]

    return {"blocks": blocks, "tables": tables}


@mcp.tool()
def write_xlsx(path: str, sheets: list[dict], font: str = "Arial") -> dict:
    """Create a formatted .xlsx file. `path` must end in .xlsx and resolve
    under /home/shanvi/claude_agent or data/documents/ in this repo.
    `font` sets the workbook-wide default font (Arial/Times New Roman —
    match the official skill's "professional font" convention).

    `sheets`: list of {
      "name": str,
      "headers": list[str]|None,       # becomes a bold, shaded header row
      "rows": list[list],              # plain data (formula-injection sanitized)
      "formulas": dict[str,str]|None,  # cell ref -> real formula, e.g. {"D2": "=SUM(B2:C2)"}
      "column_formats": dict[str,str]|None,  # column letter -> Excel number format,
                                              # e.g. {"B": "$#,##0;($#,##0);-", "C": "0.0%"}
      "cell_styles": dict[str,str]|None,     # cell ref -> one of the financial-model
                                              # roles below, applied as color/fill
      "column_widths": dict[str,int]|None,   # column letter -> width in characters
      "freeze_header": bool,           # default True if headers given
    }

    `cell_styles` roles (financial-model convention): "input" (blue text —
    hardcoded inputs/scenario levers), "formula" (black — default, rarely
    needs setting explicitly), "link" (green — reference to another sheet),
    "external_link" (red — reference to another file), "assumption"
    (yellow fill — key assumption or a cell meant for the user to fill in).

    `rows` values starting with =, +, -, or @ are automatically escaped as
    literal text to prevent formula injection; `formulas` is the ONLY way a
    real formula gets written. Note: openpyxl cannot compute formula
    results — a formula cell has no cached value until the file is opened
    once in Excel/LibreOffice, so read_xlsx will show None for it until then.
    """
    try:
        dest = _validate_write_path(path, ".xlsx")
    except ValueError as exc:
        return {"written": False, "error": str(exc)}

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    for sheet in sheets:
        ws = wb.create_sheet(title=sheet["name"][:31])
        row_offset = 1
        headers = sheet.get("headers")
        if headers:
            for col, val in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col, value=str(val))
                cell.font = Font(name=font, bold=True)
                cell.fill = PatternFill(start_color=_HEADER_FILL, end_color=_HEADER_FILL, fill_type="solid")
            row_offset = 2
            if sheet.get("freeze_header", True):
                ws.freeze_panes = "A2"

        for r, row in enumerate(sheet.get("rows", []), start=row_offset):
            for c, val in enumerate(row, start=1):
                cell = ws.cell(row=r, column=c, value=_sanitize_cell(val))
                cell.font = Font(name=font)

        for ref, formula in (sheet.get("formulas") or {}).items():
            ws[ref] = formula
            ws[ref].font = Font(name=font)

        for col_letter, fmt in (sheet.get("column_formats") or {}).items():
            for row_cells in ws[col_letter]:
                row_cells.number_format = fmt

        for ref, role in (sheet.get("cell_styles") or {}).items():
            style = _FINANCIAL_CELL_STYLES.get(role)
            if not style:
                continue
            cell = ws[ref]
            if "color" in style:
                cell.font = Font(name=font, color=style["color"], bold=cell.font.bold)
            if "fill" in style:
                cell.fill = PatternFill(start_color=style["fill"], end_color=style["fill"], fill_type="solid")

        for col_letter, width in (sheet.get("column_widths") or {}).items():
            ws.column_dimensions[col_letter].width = width

    wb.save(str(dest))
    return {"written": True, "path": str(dest)}


@mcp.tool()
def read_xlsx(path: str, sheet: str | None = None) -> dict:
    """Read an .xlsx file's cell values back out, keyed by sheet name (or
    just the requested `sheet`). Each sheet's value is a list of rows (list
    of cell values). No path restriction on reads — see read_docx docstring
    for why that's safe here. Formula cells read as None unless the file
    was opened at least once in Excel/LibreOffice since being written (see
    write_xlsx docstring).
    """
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception as exc:
        return {"error": str(exc)}

    names = [sheet] if sheet else wb.sheetnames
    result = {}
    for name in names:
        if name not in wb.sheetnames:
            return {"error": f"No sheet named {name!r}. Available: {wb.sheetnames}"}
        ws = wb[name]
        result[name] = [list(row) for row in ws.iter_rows(values_only=True)]

    return result


if __name__ == "__main__":
    mcp.run()
